import discord
from discord.utils import get
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import json
import os
import date_converter as dc

document = Document()
config = json.load(open('config.json', 'r'))\

intents = discord.Intents.none()
intents.reactions = True
intents.members = True
intents.guilds = True

client = discord.Client()
token = config['token']
ch = ''
opt = ''
cid = ''
img = 'y'

@client.event
async def on_ready():
	print(f'Logged in as {client.user}')
    
	if opt == '1':
		ch = await client.fetch_user(cid)
	elif opt == '2':
		ch = client.get_channel(int(cid))

	print('Working...')
	print('Note: This may take some time if there are a lot of messages!')
  
	messages = await ch.history(limit=20).flatten()
	messages.reverse()

	prev_user = ''
	prev_date = ''
	p = None
	h = None

	for message in messages:
		msg_date = dc.get_date(str(message.created_at).split(' ')[0])

		if(prev_date != msg_date):
			h = document.add_heading(msg_date, level=2)
			prev_date = msg_date
			prev_user = ''

		if(prev_user != message.author.name):
			if message.attachments:
				for attachment in message.attachments:
					p = document.add_paragraph(message.author.name + ':\t' + attachment.url)
			else:	
				p = document.add_paragraph(message.author.name + ':\t' + message.content)

			prev_user = message.author.name
			p.line_spacing = 1
		else:
			p = document.add_paragraph('\t' + message.content)
			p.line_spacing = 0

		paragraph_format = p.paragraph_format
		paragraph_format.tab_stops.add_tab_stop(Inches(0.6))
		paragraph_format.space_after = Pt(0)
		paragraph_format.space_before = Pt(0)

		if not os.path.exists('exported'):
			os.makedirs('exported')

		export = ''

		if opt == '1':
			export = 'exported/' + ch.name + '.docx'
		elif opt == '2':
			export = 'exported/' + cid + '.docx'

		document.save(export)
                
	print('Done!')

if __name__ == '__main__':
	opt = input('Do you want to export DM(1) or Server Channel(2)? ')

	if opt == '1':
		cid = input('Enter User ID: ')
	elif opt == '2':
		cid = input('Enter Channel ID: ')
	else:
		print('Invalid choice!')
		exit()

	#img = input('Import images? (y/n, default is yes) ')

	client.run(token, bot=False)
